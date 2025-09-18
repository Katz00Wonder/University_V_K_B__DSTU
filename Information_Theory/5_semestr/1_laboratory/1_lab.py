import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from collections import Counter
import math
import re
import PyPDF2
import docx
import os

class TextEntropyAnalyzer:
    def __init__(self, root):
        self.root = root
        self.root.title("Анализатор энтропии текста")
        self.root.geometry("1200x800")
        
        self.text_content = ""
        self.char_frequencies = {}
        self.probabilities = {}
        
        self.setup_ui()
        
    def setup_ui(self):
        # Основной фрейм
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Кнопки управления
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=0, column=0, columnspan=2, pady=10, sticky=(tk.W, tk.E))
        
        ttk.Button(button_frame, text="Загрузить файл", 
                  command=self.load_file).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Полная гистограмма", 
                  command=self.show_full_histogram).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Раздельные гистограммы", 
                  command=self.show_separate_histograms).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Топ-20 символов", 
                  command=self.show_top_histogram).pack(side=tk.LEFT, padx=5)
        
  
        self.file_info_label = ttk.Label(main_frame, text="Файл: не выбран")
        self.file_info_label.grid(row=1, column=0, sticky=tk.W, pady=5)
        

        info_frame = ttk.Frame(main_frame)
        info_frame.grid(row=2, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=10)
        

        self.entropy_label = ttk.Label(info_frame, text="Энтропия: не рассчитана")
        self.entropy_label.pack(anchor=tk.W)
        
        self.total_chars_label = ttk.Label(info_frame, text="Всего символов: 0")
        self.total_chars_label.pack(anchor=tk.W)
        
        self.unique_chars_label = ttk.Label(info_frame, text="Уникальных символов: 0")
        self.unique_chars_label.pack(anchor=tk.W)
        

        self.figure_frame = ttk.Frame(main_frame)
        self.figure_frame.grid(row=3, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=10)
        

        table_frame = ttk.Frame(main_frame)
        table_frame.grid(row=2, column=1, rowspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10)
        

        columns = ('symbol', 'count', 'frequency', 'probability')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=20)
        
        self.tree.heading('symbol', text='Символ')
        self.tree.heading('count', text='Количество')
        self.tree.heading('frequency', text='Частота')
        self.tree.heading('probability', text='Вероятность')
        
        self.tree.column('symbol', width=80)
        self.tree.column('count', width=100)
        self.tree.column('frequency', width=100)
        self.tree.column('probability', width=100)
        

        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)
        
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        

        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(3, weight=1)
        
    def load_file(self):
        file_path = filedialog.askopenfilename(
            title="Выберите файл",
            filetypes=[
                ("Текстовые файлы", "*.txt"),
                ("PDF файлы", "*.pdf"),
                ("Word документы", "*.docx"),
                ("Word документы", "*.doc"),
                ("Все файлы", "*.*")
            ]
        )
        
        if file_path:
            try:
                self.file_info_label.config(text=f"Файл: {os.path.basename(file_path)}")
                file_extension = os.path.splitext(file_path)[1].lower()
                
                if file_extension == '.pdf':
                    self.text_content = self.read_pdf(file_path)
                elif file_extension in ['.docx', '.doc']:
                    self.text_content = self.read_word(file_path)
                else:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        self.text_content = file.read()
                
                self.analyze_text()
                messagebox.showinfo("Успех", "Файл успешно загружен и проанализирован")
                
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось загрузить файл: {str(e)}")
    
    def read_pdf(self, file_path):
        """Чтение текста из PDF файла"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise Exception(f"Ошибка чтения PDF: {str(e)}")
    
    def read_word(self, file_path):
        """Чтение текста из Word документа"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Также читаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Ошибка чтения Word документа: {str(e)}")
    
    def clean_text(self, text):
        """Очистка текста от указанных символов"""
        if not text:
            return ""
            
        # Удаляем указанные символы
        chars_to_remove = r'[@#$^&*{}[\]<>/\\|=+`]'
        cleaned_text = re.sub(chars_to_remove, '', text)
        return cleaned_text
    
    def analyze_text(self):
        # Очищаем текст
        cleaned_text = self.clean_text(self.text_content)
        
        if not cleaned_text:
            messagebox.showwarning("Предупреждение", "Текст после очистки пуст")
            return
        
        # Подсчитываем частоты символов
        total_chars = len(cleaned_text)
        char_counter = Counter(cleaned_text)
        
        self.char_frequencies = dict(char_counter)
        self.probabilities = {char: count/total_chars for char, count in char_counter.items()}
        
        # Вычисляем энтропию
        entropy = self.calculate_entropy()
        
        # Обновляем информацию
        self.entropy_label.config(text=f"Энтропия: {entropy:.4f}")
        self.total_chars_label.config(text=f"Всего символов: {total_chars}")
        self.unique_chars_label.config(text=f"Уникальных символов: {len(self.char_frequencies)}")
        
        # Обновляем таблицу
        self.update_table()
    
    def calculate_entropy(self):
        entropy = 0.0
        for prob in self.probabilities.values():
            if prob > 0:
                entropy -= prob * math.log2(prob)
        return entropy
    
    def update_table(self):
        # Очищаем таблицу
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # Сортируем символы по частоте
        sorted_chars = sorted(self.probabilities.items(), key=lambda x: x[1], reverse=True)
        
        # Заполняем таблицу
        for char, prob in sorted_chars:
            count = self.char_frequencies[char]
            # Экранируем специальные символы для отображения
            display_char = self.escape_char(char)
            self.tree.insert('', 'end', values=(
                display_char, 
                count, 
                f"{count/len(self.text_content)*100:.2f}%",
                f"{prob:.6f}"
            ))
    
    def escape_char(self, char):
        # Экранирование специальных символов для отображения
        if char == ' ':
            return "' ' (пробел)"
        elif char == '\t':
            return "\\t (табуляция)"
        elif char == '\n':
            return "\\n (новая строка)"
        elif char == '\r':
            return "\\r (возврат каретки)"
        else:
            return char
    
    def show_full_histogram(self):
        if not self.probabilities:
            messagebox.showwarning("Предупреждение", "Сначала загрузите файл")
            return
        
        self.create_histogram(list(self.probabilities.keys()), list(self.probabilities.values()), 
                             "Полная гистограмма всех символов")
    
    def show_separate_histograms(self):
        if not self.probabilities:
            messagebox.showwarning("Предупреждение", "Сначала загрузите файл")
            return
        
        # Разделяем символы по категориям
        cyrillic_chars = {}
        latin_chars = {}
        digit_chars = {}
        punctuation_chars = {}
        other_chars = {}
        
        for char, prob in self.probabilities.items():
            if '\u0400' <= char <= '\u04FF':
                cyrillic_chars[char] = prob
            elif 'a' <= char <= 'z' or 'A' <= char <= 'Z':  
                latin_chars[char] = prob
            elif '0' <= char <= '9': 
                digit_chars[char] = prob
            elif char in '.,!?;:"\'()-': 
                punctuation_chars[char] = prob
            else:
                other_chars[char] = prob
        
        
        fig, axes = plt.subplots(2, 2, figsize=(15, 10))
        fig.suptitle('Раздельные гистограммы символов')
        

        if cyrillic_chars:
            self.plot_on_axis(axes[0, 0], cyrillic_chars, 'Кириллица')
        else:
            axes[0, 0].set_title('Кириллица (нет данных)')
        

        if latin_chars:
            self.plot_on_axis(axes[0, 1], latin_chars, 'Латиница')
        else:
            axes[0, 1].set_title('Латиница (нет данных)')
        
        
        if digit_chars or punctuation_chars:
            combined = {**digit_chars, **punctuation_chars}
            self.plot_on_axis(axes[1, 0], combined, 'Цифры и знаки препинания')
        else:
            axes[1, 0].set_title('Цифры и пунктуация (нет данных)')
        
    
        if other_chars:
            self.plot_on_axis(axes[1, 1], other_chars, 'Прочие символы')
        else:
            axes[1, 1].set_title('Прочие символы (нет данных)')
        
        plt.tight_layout()
        self.display_figure(fig)
    
    def show_top_histogram(self):
        if not self.probabilities:
            messagebox.showwarning("Предупреждение", "Сначала загрузите файл")
            return
        
        
        sorted_probs = sorted(self.probabilities.items(), key=lambda x: x[1], reverse=True)[:20]
        top_chars = [self.escape_char(char) for char, _ in sorted_probs]
        top_probs = [prob for _, prob in sorted_probs]
        
        self.create_histogram(top_chars, top_probs, "Топ-20 самых частых символов")
    
    def plot_on_axis(self, ax, data, title):
        chars = list(data.keys())
        probs = list(data.values())
        

        display_chars = [self.escape_char(char) for char in chars]
        
        ax.bar(range(len(chars)), probs)
        ax.set_title(title)
        ax.set_ylabel('Вероятность')
        ax.set_xticks(range(len(chars)))
        ax.set_xticklabels(display_chars, rotation=45, ha='right')
    
    def create_histogram(self, chars, probs, title):
        fig, ax = plt.subplots(figsize=(12, 6))
        
 
        display_chars = [self.escape_char(char) for char in chars]
        
        ax.bar(range(len(chars)), probs)
        ax.set_title(title)
        ax.set_xlabel('Символы')
        ax.set_ylabel('Вероятность')
        ax.set_xticks(range(len(chars)))
        ax.set_xticklabels(display_chars, rotation=45, ha='right')
        
        plt.tight_layout()
        self.display_figure(fig)
    
    def display_figure(self, fig):

        for widget in self.figure_frame.winfo_children():
            widget.destroy()
        

        canvas = FigureCanvasTkAgg(fig, master=self.figure_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

def main():
    root = tk.Tk()
    app = TextEntropyAnalyzer(root)
    root.mainloop()

if __name__ == "__main__":
    main()