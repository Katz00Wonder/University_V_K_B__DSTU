import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from collections import Counter
import math
import re
import PyPDF2
import docx
import os

class TextEntropyAnalyzer:
    def __init__(self, root):
        self.root = root
        self.root.title("Анализатор энтропии файлов и текста")
        self.root.geometry("1400x900")
        
        self.text_content = ""
        self.char_frequencies = {}
        self.probabilities = {}
        self.file_data = b""
        
        self.setup_ui()
        
    def setup_ui(self):
        # Основной фрейм
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Кнопки управления
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=0, column=0, columnspan=2, pady=10, sticky=(tk.W, tk.E))
        
        ttk.Button(button_frame, text="Загрузить файл", 
                  command=self.load_file).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Гистограмма байтов файла", 
                  command=self.show_file_bytes_histogram).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Полная гистограмма символов", 
                  command=self.show_full_histogram).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Раздельные гистограммы", 
                  command=self.show_separate_histograms).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Топ-20 символов", 
                  command=self.show_top_histogram).pack(side=tk.LEFT, padx=5)
        
        # Информация о файле
        self.file_info_label = ttk.Label(main_frame, text="Файл: не выбран")
        self.file_info_label.grid(row=1, column=0, sticky=tk.W, pady=5)
        
        # Фрейм для информации об энтропии
        entropy_info_frame = ttk.LabelFrame(main_frame, text="Результаты энтропии", padding="10")
        entropy_info_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=10)
        
        # Энтропия файла
        self.file_entropy_label = ttk.Label(entropy_info_frame, text="Энтропия файла: не рассчитана")
        self.file_entropy_label.pack(anchor=tk.W)
        
        # Энтропия текста
        self.text_entropy_label = ttk.Label(entropy_info_frame, text="Энтропия текста: не рассчитана")
        self.text_entropy_label.pack(anchor=tk.W)
        
        # Сравнение энтропий
        self.comparison_label = ttk.Label(entropy_info_frame, text="Сравнение: не доступно")
        self.comparison_label.pack(anchor=tk.W)
        
        # Информация о файле
        file_info_frame = ttk.LabelFrame(main_frame, text="Информация о файле", padding="10")
        file_info_frame.grid(row=2, column=1, sticky=(tk.W, tk.E), pady=10, padx=10)
        
        self.file_size_label = ttk.Label(file_info_frame, text="Размер файла: 0 байт")
        self.file_size_label.pack(anchor=tk.W)
        
        self.total_chars_label = ttk.Label(file_info_frame, text="Всего символов текста: 0")
        self.total_chars_label.pack(anchor=tk.W)
        
        self.unique_chars_label = ttk.Label(file_info_frame, text="Уникальных символов: 0")
        self.unique_chars_label.pack(anchor=tk.W)
        
        # Фрейм для графиков
        self.figure_frame = ttk.Frame(main_frame)
        self.figure_frame.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=10)
        
        # Таблица символов
        table_frame = ttk.LabelFrame(main_frame, text="Статистика символов", padding="10")
        table_frame.grid(row=4, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=10)
        
        # Создание таблицы
        columns = ('symbol', 'count', 'frequency', 'probability')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=15)
        
        self.tree.heading('symbol', text='Символ')
        self.tree.heading('count', text='Количество')
        self.tree.heading('frequency', text='Частота')
        self.tree.heading('probability', text='Вероятность')
        
        self.tree.column('symbol', width=100)
        self.tree.column('count', width=100)
        self.tree.column('frequency', width=100)
        self.tree.column('probability', width=120)
        
        # Скроллбар для таблиции
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)
        
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Настройка весов для растягивания
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(3, weight=1)
        main_frame.rowconfigure(4, weight=1)
        
    def load_file(self):
        file_path = filedialog.askopenfilename(
            title="Выберите файл",
            filetypes=[
                ("Текстовые файлы", "*.txt"),
                ("PDF файлы", "*.pdf"),
                ("Word документы", "*.docx"),
                ("Word документы", "*.doc"),
                ("Все файлы", "*.*")
            ]
        )
        
        if file_path:
            try:
                self.file_info_label.config(text=f"Файл: {os.path.basename(file_path)}")
                
                
                with open(file_path, 'rb') as file:
                    self.file_data = file.read()
                
                
                file_entropy = self.calculate_file_entropy()
                self.file_entropy_label.config(text=f"Энтропия файла: {file_entropy:.4f} бит/байт")
                
                
                file_size = len(self.file_data)
                self.file_size_label.config(text=f"Размер файла: {file_size} байт")
                
                
                file_extension = os.path.splitext(file_path)[1].lower()
                text_entropy = None
                
                if file_extension == '.pdf':
                    self.text_content = self.read_pdf(file_path)
                elif file_extension in ['.docx', '.doc']:
                    self.text_content = self.read_word(file_path)
                else:
                   
                    encodings = ['utf-8', 'cp1251', 'latin-1', 'ascii']
                    for encoding in encodings:
                        try:
                            with open(file_path, 'r', encoding=encoding) as file:
                                self.text_content = file.read()
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        
                        try:
                            self.text_content = self.file_data.decode('utf-8', errors='ignore')
                        except:
                            self.text_content = ""
                
                
                if self.text_content and len(self.text_content.strip()) > 0:
                    self.analyze_text()
                    text_entropy = self.calculate_text_entropy()
                    self.text_entropy_label.config(text=f"Энтропия текста: {text_entropy:.4f} бит/символ")
                    
                    
                    if text_entropy is not None:
                        diff = abs(file_entropy - text_entropy)
                        comparison_text = f"Разница энтропий: {diff:.4f} бит"
                        if file_entropy > text_entropy:
                            comparison_text += " (файл имеет большую энтропию)"
                        elif text_entropy > file_entropy:
                            comparison_text += " (текст имеет большую энтропию)"
                        else:
                            comparison_text += " (энтропии равны)"
                        self.comparison_label.config(text=comparison_text)
                else:
                    self.text_entropy_label.config(text="Энтропия текста: файл не содержит текст")
                    self.comparison_label.config(text="Сравнение: не доступно (нет текста)")
                    self.text_content = ""
                
                messagebox.showinfo("Успех", "Файл успешно загружен и проанализирован")
                
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось загрузить файл: {str(e)}")
    
    def calculate_file_entropy(self):
        """Расчет энтропии Шеннона для файла (по байтам)"""
        if not self.file_data:
            return 0
        
        byte_counts = Counter(self.file_data)
        total_bytes = len(self.file_data)
        entropy = 0
        
        for count in byte_counts.values():
            probability = count / total_bytes
            if probability > 0:
                entropy -= probability * math.log2(probability)
        
        return entropy
    
    def calculate_text_entropy(self):
        """Расчет энтропии Шеннона для текста"""
        if not self.probabilities:
            return 0
        
        entropy = 0.0
        for prob in self.probabilities.values():
            if prob > 0:
                entropy -= prob * math.log2(prob)
        return entropy
    
    def read_pdf(self, file_path):
        """Чтение текста из PDF файла"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise Exception(f"Ошибка чтения PDF: {str(e)}")
    
    def read_word(self, file_path):
        """Чтение текста из Word документа"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Ошибка чтения Word документа: {str(e)}")
    
    def clean_text(self, text):
        """Очистка текста от указанных символов"""
        if not text:
            return ""
            
        
        chars_to_remove = r'[@#$^&*{}[\]<>/\\|=+`]'
        cleaned_text = re.sub(chars_to_remove, '', text)
        return cleaned_text
    
    def analyze_text(self):
        
        cleaned_text = self.clean_text(self.text_content)
        
        if not cleaned_text:
            messagebox.showwarning("Предупреждение", "Текст после очистки пуст")
            return
        
        
        total_chars = len(cleaned_text)
        char_counter = Counter(cleaned_text)
        
        self.char_frequencies = dict(char_counter)
        self.probabilities = {char: count/total_chars for char, count in char_counter.items()}
        
        
        self.total_chars_label.config(text=f"Всего символов текста: {total_chars}")
        self.unique_chars_label.config(text=f"Уникальных символов: {len(self.char_frequencies)}")
        
        
        self.update_table()
    
    def update_table(self):
        
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        
        sorted_chars = sorted(self.probabilities.items(), key=lambda x: x[1], reverse=True)
        
        
        for char, prob in sorted_chars:
            count = self.char_frequencies[char]
            
            display_char = self.escape_char(char)
            self.tree.insert('', 'end', values=(
                display_char, 
                count, 
                f"{count/len(self.text_content)*100:.2f}%",
                f"{prob:.6f}"
            ))
    
    def escape_char(self, char):
        
        if char == ' ':
            return "' ' (пробел)"
        elif char == '\t':
            return "\\t (табуляция)"
        elif char == '\n':
            return "\\n (новая строка)"
        elif char == '\r':
            return "\\r (возврат каретки)"
        else:
            return char
    
    def show_file_bytes_histogram(self):
        """Гистограмма распределения байтов в файле"""
        if not self.file_data:
            messagebox.showwarning("Предупреждение", "Сначала загрузите файл")
            return
        
        
        byte_counts = Counter(self.file_data)
        bytes_list = list(range(256))  
        frequencies = [byte_counts.get(byte, 0) for byte in bytes_list]
        
        
        total_bytes = len(self.file_data)
        if total_bytes > 0:
            frequencies = [freq / total_bytes for freq in frequencies]
        
        
        fig, ax = plt.subplots(figsize=(16, 8))
        
        
        bars = ax.bar(bytes_list, frequencies, alpha=0.7, color='skyblue', edgecolor='black', linewidth=0.3)
        ax.set_xlabel('Байт (десятичное значение)')
        ax.set_ylabel('Частота появления')
        ax.set_title('Распределение байтов в файле (0-255)')
        ax.set_xlim(-5, 260)
        
       
        ax.set_xticks(range(0, 256, 16)) 
        ax.grid(True, alpha=0.3, axis='y')
        
        
        if total_bytes > 0:
            
            top_bytes = sorted([(byte, freq) for byte, freq in enumerate(frequencies)], 
                             key=lambda x: x[1], reverse=True)[:5]
            
            info_text = "Самые частые байты:\n"
            for i, (byte, freq) in enumerate(top_bytes):
                if freq > 0:
                    char_desc = self.get_byte_description(byte)
                    info_text += f"{byte} (0x{byte:02X}) - {freq:.4f} {char_desc}\n"
            
            
            ax.text(0.02, 0.98, info_text, transform=ax.transAxes, fontsize=10,
                   verticalalignment='top', bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.8))
        
        
        significant_bytes = [0, 9, 10, 13, 32, 127, 255]
        for byte in significant_bytes:
            if frequencies[byte] > 0.001:  
                ax.text(byte, frequencies[byte] + 0.002, str(byte), 
                       ha='center', va='bottom', fontsize=8, fontweight='bold')
        
        plt.tight_layout()
        self.display_figure(fig)
    
    def get_byte_description(self, byte):
        """Получить текстовое описание байта"""
        if byte == 0:
            return "(NUL)"
        elif byte == 9:
            return "(TAB)"
        elif byte == 10:
            return "(LF)"
        elif byte == 13:
            return "(CR)"
        elif byte == 32:
            return "(SPACE)"
        elif byte == 127:
            return "(DEL)"
        elif 32 <= byte <= 126:
            return f"('{chr(byte)}')"
        else:
            return ""
    
    def show_full_histogram(self):
        if not self.probabilities:
            messagebox.showwarning("Предупреждение", "Сначала загрузите файл с текстом")
            return
        
        self.create_histogram(list(self.probabilities.keys()), list(self.probabilities.values()), 
                             "Полная гистограмма всех символов текста")
    
    def show_separate_histograms(self):
        if not self.probabilities:
            messagebox.showwarning("Предупреждение", "Сначала загрузите файл с текстом")
            return
        
        # Разделяем символы по категориям
        cyrillic_chars = {}
        latin_chars = {}
        digit_chars = {}
        punctuation_chars = {}
        other_chars = {}
        
        for char, prob in self.probabilities.items():
            if '\u0400' <= char <= '\u04FF':
                cyrillic_chars[char] = prob
            elif 'a' <= char <= 'z' or 'A' <= char <= 'Z':  
                latin_chars[char] = prob
            elif '0' <= char <= '9': 
                digit_chars[char] = prob
            elif char in '.,!?;:"\'()-': 
                punctuation_chars[char] = prob
            else:
                other_chars[char] = prob
        
        fig, axes = plt.subplots(2, 2, figsize=(15, 10))
        fig.suptitle('Раздельные гистограммы символов текста')
        
        if cyrillic_chars:
            self.plot_on_axis(axes[0, 0], cyrillic_chars, 'Кириллица')
        else:
            axes[0, 0].set_title('Кириллица (нет данных)')
        
        if latin_chars:
            self.plot_on_axis(axes[0, 1], latin_chars, 'Латиница')
        else:
            axes[0, 1].set_title('Латиница (нет данных)')
        
        if digit_chars or punctuation_chars:
            combined = {**digit_chars, **punctuation_chars}
            self.plot_on_axis(axes[1, 0], combined, 'Цифры и знаки препинания')
        else:
            axes[1, 0].set_title('Цифры и пунктуация (нет данных)')
        
        if other_chars:
            self.plot_on_axis(axes[1, 1], other_chars, 'Прочие символы')
        else:
            axes[1, 1].set_title('Прочие символы (нет данных)')
        
        plt.tight_layout()
        self.display_figure(fig)
    
    def show_top_histogram(self):
        if not self.probabilities:
            messagebox.showwarning("Предупреждение", "Сначала загрузите файл с текстом")
            return
        
        sorted_probs = sorted(self.probabilities.items(), key=lambda x: x[1], reverse=True)[:20]
        top_chars = [self.escape_char(char) for char, _ in sorted_probs]
        top_probs = [prob for _, prob in sorted_probs]
        
        self.create_histogram(top_chars, top_probs, "Топ-20 самых частых символов текста")
    
    def plot_on_axis(self, ax, data, title):
        chars = list(data.keys())
        probs = list(data.values())
        
        display_chars = [self.escape_char(char) for char in chars]
        
        ax.bar(range(len(chars)), probs)
        ax.set_title(title)
        ax.set_ylabel('Вероятность')
        ax.set_xticks(range(len(chars)))
        ax.set_xticklabels(display_chars, rotation=45, ha='right')
    
    def create_histogram(self, chars, probs, title):
        fig, ax = plt.subplots(figsize=(12, 6))
        
        display_chars = [self.escape_char(char) for char in chars]
        
        ax.bar(range(len(chars)), probs)
        ax.set_title(title)
        ax.set_xlabel('Символы')
        ax.set_ylabel('Вероятность')
        ax.set_xticks(range(len(chars)))
        ax.set_xticklabels(display_chars, rotation=45, ha='right')
        
        plt.tight_layout()
        self.display_figure(fig)
    
    def display_figure(self, fig):
        # Очищаем фрейм от предыдущих графиков
        for widget in self.figure_frame.winfo_children():
            widget.destroy()
        
        # Создаем новый canvas для графика
        canvas = FigureCanvasTkAgg(fig, master=self.figure_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

def main():
    root = tk.Tk()
    app = TextEntropyAnalyzer(root)
    root.mainloop()

if __name__ == "__main__":
    main()